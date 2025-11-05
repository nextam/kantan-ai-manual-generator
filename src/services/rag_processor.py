"""
File: rag_processor.py
Purpose: RAG processing pipeline for reference materials
Main functionality: Text extraction, chunking, metadata extraction, embedding generation
Dependencies: PyPDF2, pdfplumber, python-docx, openpyxl, google.genai
"""

import os
import re
import tempfile
from typing import List, Dict, Any, Optional, Tuple
import json

# Text extraction libraries
import PyPDF2
import pdfplumber
from docx import Document
from openpyxl import load_workbook

# Google Gemini for embeddings and metadata
from google import genai
from google.genai import types

from src.infrastructure.s3_manager import s3_manager


class RAGProcessor:
    """
    RAG processing pipeline for reference materials
    
    Pipeline:
    1. Download file from S3
    2. Extract text based on file type
    3. Extract metadata using Gemini
    4. Chunk text into manageable pieces
    5. Generate vector embeddings
    6. Index in ElasticSearch
    """
    
    def __init__(self):
        self.project_id = os.getenv('PROJECT_ID', 'kantan-ai-database')
        self.location = os.getenv('VERTEX_AI_LOCATION', 'us-central1')
        
        # Initialize Gemini client
        self.client = genai.Client(
            vertexai=True,
            project=self.project_id,
            location=self.location
        )
        
        # Chunking configuration
        self.chunk_size = 1000  # Target tokens per chunk
        self.chunk_overlap = 50  # Overlap tokens
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file
        
        Uses pdfplumber for better text extraction quality
        Falls back to PyPDF2 if pdfplumber fails
        
        Args:
            file_path: Local path to PDF file
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            text_parts = []
            metadata = {'pages': 0, 'method': 'pdfplumber'}
            
            # Try pdfplumber first (better quality)
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            if text_parts:
                return '\n\n'.join(text_parts), metadata
        
        except Exception as e:
            print(f"pdfplumber failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        try:
            text_parts = []
            metadata = {'pages': 0, 'method': 'PyPDF2'}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            return '\n\n'.join(text_parts), metadata
        
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from Word document
        
        Args:
            file_path: Local path to DOCX file
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = Document(file_path)
            
            text_parts = []
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)
                
                if table_text:
                    text_parts.append('[Table]\n' + '\n'.join(table_text))
            
            return '\n\n'.join(text_parts), metadata
        
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_xlsx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from Excel file
        
        Args:
            file_path: Local path to XLSX file
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            text_parts = []
            metadata = {'sheets': len(workbook.sheetnames)}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                sheet_text = [f"[Sheet: {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                if len(sheet_text) > 1:
                    text_parts.append('\n'.join(sheet_text))
            
            return '\n\n'.join(text_parts), metadata
        
        except Exception as e:
            raise Exception(f"Failed to extract XLSX text: {str(e)}")
    
    def extract_text(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from file based on type
        
        Args:
            file_path: Local path to file
            file_type: File type (pdf, docx, xlsx, csv)
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_type == 'xlsx':
            return self.extract_text_from_xlsx(file_path)
        elif file_type == 'csv':
            # Simple CSV reading
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text, {'type': 'csv'}
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    def extract_metadata_with_gemini(self, text: str, title: str) -> Dict[str, Any]:
        """
        Extract metadata from text using Gemini
        
        Extracts:
        - Document type/category
        - Key topics
        - Summary
        
        Args:
            text: Extracted text
            title: Material title
        
        Returns:
            Metadata dictionary
        """
        try:
            # Truncate text if too long
            max_text_length = 10000
            truncated_text = text[:max_text_length] if len(text) > max_text_length else text
            
            prompt = f"""Analyze this reference material and extract metadata.

Title: {title}

Content:
{truncated_text}

Please provide:
1. Document type/category (e.g., manual, procedure, specification, guideline)
2. Key topics (3-5 main topics)
3. Brief summary (1-2 sentences)

Format your response as JSON:
{{
    "document_type": "...",
    "key_topics": ["topic1", "topic2", "topic3"],
    "summary": "..."
}}
"""
            
            response = self.client.models.generate_content(
                model='gemini-2.0-flash-exp',
                contents=prompt
            )
            
            # Extract JSON from response
            response_text = response.text.strip()
            
            # Try to parse JSON
            try:
                # Remove markdown code blocks if present
                if '```json' in response_text:
                    response_text = response_text.split('```json')[1].split('```')[0].strip()
                elif '```' in response_text:
                    response_text = response_text.split('```')[1].split('```')[0].strip()
                
                metadata = json.loads(response_text)
                return metadata
            
            except json.JSONDecodeError:
                # Fallback to basic metadata
                return {
                    'document_type': 'unknown',
                    'key_topics': [],
                    'summary': 'Metadata extraction failed',
                    'raw_response': response_text
                }
        
        except Exception as e:
            print(f"Gemini metadata extraction failed: {e}")
            return {
                'document_type': 'unknown',
                'key_topics': [],
                'summary': 'Metadata extraction failed',
                'error': str(e)
            }
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks with overlap
        
        Strategy:
        1. Split by paragraphs (double newlines)
        2. Combine paragraphs until reaching target chunk size
        3. Add overlap from previous chunk
        
        Args:
            text: Text to chunk
            chunk_size: Target tokens per chunk (default: self.chunk_size)
            overlap: Overlap tokens (default: self.chunk_overlap)
        
        Returns:
            List of chunks with metadata
        """
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap
        
        # Simple token estimation: 1 token ≈ 4 characters
        char_per_token = 4
        target_chars = chunk_size * char_per_token
        overlap_chars = overlap * char_per_token
        
        # Split by paragraphs
        paragraphs = re.split(r'\n\n+', text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_length = len(para)
            
            # If adding this paragraph exceeds target, save current chunk
            if current_length + para_length > target_chars and current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'char_length': len(chunk_text),
                    'estimated_tokens': len(chunk_text) // char_per_token
                })
                
                # Start new chunk with overlap
                overlap_text = chunk_text[-overlap_chars:] if len(chunk_text) > overlap_chars else chunk_text
                current_chunk = [overlap_text, para]
                current_length = len(overlap_text) + para_length
            else:
                current_chunk.append(para)
                current_length += para_length
        
        # Add final chunk
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            chunks.append({
                'text': chunk_text,
                'char_length': len(chunk_text),
                'estimated_tokens': len(chunk_text) // char_per_token
            })
        
        return chunks
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings using Vertex AI text-embedding-004
        
        Args:
            texts: List of texts to embed
        
        Returns:
            List of embedding vectors (768-dim)
        """
        try:
            embeddings = []
            
            # Process in batches of 100
            batch_size = 100
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                # Use Vertex AI embeddings
                response = self.client.models.embed_content(
                    model='text-embedding-004',
                    contents=batch
                )
                
                for embedding in response.embeddings:
                    embeddings.append(embedding.values)
            
            return embeddings
        
        except Exception as e:
            raise Exception(f"Embedding generation failed: {str(e)}")
    
    def process_material(self, material_id: int, company_id: int, file_path_s3: str,
                        file_type: str, title: str) -> Dict[str, Any]:
        """
        Complete RAG processing pipeline
        
        Args:
            material_id: ReferenceMaterial.id
            company_id: Company.id
            file_path_s3: S3 URI
            file_type: File type
            title: Material title
        
        Returns:
            Processing results
        """
        temp_file = None
        
        try:
            # Step 1: Download from S3
            s3_key = file_path_s3.replace(f's3://{s3_manager.bucket_name}/', '')
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as tf:
                temp_file = tf.name
                s3_manager.download_file(s3_key, temp_file)
            
            # Step 2: Extract text
            extracted_text, extraction_metadata = self.extract_text(temp_file, file_type)
            
            if not extracted_text or len(extracted_text) < 10:
                raise Exception("Insufficient text extracted from file")
            
            # Step 3: Extract metadata with Gemini
            gemini_metadata = self.extract_metadata_with_gemini(extracted_text, title)
            
            # Step 4: Chunk text
            chunks = self.chunk_text(extracted_text)
            
            if not chunks:
                raise Exception("No chunks generated from text")
            
            # Step 5: Generate embeddings
            chunk_texts = [chunk['text'] for chunk in chunks]
            embeddings = self.generate_embeddings(chunk_texts)
            
            # Combine chunks with embeddings
            processed_chunks = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                processed_chunks.append({
                    'chunk_index': i,
                    'text': chunk['text'],
                    'char_length': chunk['char_length'],
                    'estimated_tokens': chunk['estimated_tokens'],
                    'embedding': embedding,
                    'metadata': {}
                })
            
            return {
                'success': True,
                'extracted_text_length': len(extracted_text),
                'extraction_metadata': extraction_metadata,
                'gemini_metadata': gemini_metadata,
                'chunk_count': len(processed_chunks),
                'chunks': processed_chunks
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
        
        finally:
            # Clean up temp file
            if temp_file and os.path.exists(temp_file):
                try:
                    os.unlink(temp_file)
                except Exception as e:
                    print(f"Failed to delete temp file: {e}")


rag_processor = RAGProcessor()
